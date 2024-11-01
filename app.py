#!/usr/bin/env python
# coding: utf-8

# <h1>1. Set-up</h1>

# In[145]:

import streamlit as st
import requests
import re
import os
import yaml
import pandas as pd
import json
from bs4 import BeautifulSoup
from newspaper import Article
from langchain.document_loaders import PyPDFLoader


# In[146]:


from openai import OpenAI
from getpass import getpass

# Set up and run this Streamlit App
import streamlit as st

# region <--------- Streamlit App Configuration --------->
st.set_page_config(
    layout="centered",
    page_title="AI Cyber Threat Intelligence App"
)
# endregion <--------- Streamlit App Configuration --------->

# App Title
st.title("AI Cyber Threat Intelligence App")

# App Description
st.markdown("""
Welcome to the AI Cyber Threat Intelligence App! 

This tool is designed to streamline cyber threat analysis by taking in a news source, extracting relevant information, and translating it into actionable insights aligned with the MITRE ATT&CK framework. Here’s what it does:

1. **MITRE ATT&CK Mapping**: The app parses your news input to identify attacker tactics and techniques, which it categorizes by MITRE ATT&CK tactic and technique IDs.
2. **Mitigation Measures**: It provides recommended mitigation measures for each identified threat technique, offering guidance on defensive steps.
3. **Output Files**:
   - **MITRE Navigator JSON**: A `.json` file for visualizing your threat landscape in MITRE ATT&CK Navigator.
   - **KQL Query**: Custom KQL queries for use in Microsoft Sentinel, targeting the detected techniques.
   - **Comprehensive Report**: A report summarizing the threat source, including helpful, actionable insights and recommended mitigations.

### Start Threat Intelligence Process

Please enter a valid news source to begin the threat intelligence process. You can either:

1. **Enter a URL** of the threat intelligence source.
2. **Upload a PDF** file directly from your local device.

Select your preferred input method below.
""")

openai_key = st.secrets["OPENAPI_API_KEY"]
client = OpenAI(api_key=openai_key)


# ### 1.1 Only run for installation of missing packages

# In[147]:


# !pip install openai
# !pip install newspaper3k
# !pip install tiktoken
# !pip install langchain
# !pip install faiss-cpu
# !pip install -U langchain-community
# !pip install pypdf
# !pip install pyarrow==15.0.0
# !pip install datasets
# !pip install ragas
# !pip install python-docx


# ### 1.2 Only run git clone once a week to refresh splunk query data

# In[148]:


###!git clone https://github.com/splunk/security_content.git


# <h1>2. Helper Function</h1>

# In[149]:


def get_completion(prompt, model="gpt-4o-mini"):
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,  # Control randomness
    )
    return response.choices[0].message.content


# <h1>3. Content Retrieval</h1>

# ### Functions of scraping URL and PDF

# In[150]:


def scrape_article(url):
    article = Article(url)
    try:
        article.download()
        article.parse()
    except Exception as e:
        print(f"An error occurred: {e}")
        return None  # Return None if an error occurs

    # Store the scraped text in a variable
    return article.text  # Return the text instead of printing it


def load_pdf(file_name):
    try:
        loader = PyPDFLoader(file_name)
        pages = loader.load()
        pdf_text = ''.join([page.page_content for page in pages])  # Concatenate all page content
        return pdf_text
    except Exception as e:
        print(f"An error occurred while reading the PDF: {e}")
        return None


# ### User could either enter a website URL or provide a document

# In[151]:


def scrape_article(url):
    article = Article(url)
    try:
        article.download()
        article.parse()
    except Exception as e:
        print(f"An error occurred: {e}")
        return None  # Return None if an error occurs

    # Store the scraped text in a variable
    return article.text  # Return the text instead of printing it


def load_pdf(uploaded_file):
    try:
        loader = PyPDFLoader(uploaded_file)
        pages = loader.load()
        pdf_text = ''.join([page.page_content for page in pages])  # Concatenate all page content
        return pdf_text
    except Exception as e:
        print(f"An error occurred while reading the PDF: {e}")
        return None


# ### User can either enter a website URL or upload a PDF document

def user_choice():
    # Initialize scraped_text as None to ensure it's defined
    scraped_text = None

    # Radio button for selecting input method
    choice = st.radio("Choose the input method:", ["URL", "PDF"])

    if choice == "URL":
        # Ask for the URL and scrape the article
        url = st.text_input("Enter the URL of the website:")
        if url:
            scraped_text = scrape_article(url)
            if scraped_text:
                st.write("Text extracted successfully from the URL.")
            else:
                st.write("Failed to scrape the article.")

    elif choice == "PDF":
        # File uploader for PDF input
        uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
        if uploaded_file:
            scraped_text = load_pdf(uploaded_file)
            if scraped_text:
                st.write("Text extracted successfully from the PDF.")
            else:
                st.write("Failed to extract text from the PDF.")

    return scraped_text  # Ensure scraped_text is returned


# Call the function to prompt user interaction
scraped_text = user_choice()

from datetime import datetime
import pytz

singapore_tz = pytz.timezone('Asia/Singapore')
now_sg = datetime.now(singapore_tz)
export_time = now_sg.strftime('%a %b %d %Y %H:%M:%S GMT+0800 (Singapore Standard Time)')

st.markdown("""
### News Summary
""")

# Output the value
st.write(export_time)


# #### 2. Initialise report name and text summary

# In[180]:


prompt_report = f"""
News Text:
{scraped_text}

<Instructions>
You are a cybersecurity expert. Based on the following news text, extract two things:
1. Provide a news report name summarizing the key point in 1 sentence without mentioning of any MITRE ATT&CK Technique IDs but instead state the threat actors wherever available. For example: "China-Linked CeranaKeeper Targeting Southeast Asia with Data Exfiltration."
2. Write a brief summary (3-4 sentences) explaining what the news report is about, so the reader can understand it quickly.
</Instructions>

Use the following format:
### Name:
<Your name goes here>

### Summary:
<Your summary goes here>
"""


response_report = get_completion(prompt_report)


# In[181]:


# Split the output based on '###' to extract name and summary
split_output = response_report.split("###")

# Extract the title and summary
report_name = split_output[1].replace("Name:", "").strip() 
report_summary = split_output[2].replace("Summary:", "").strip()

st.write("Report Name:", report_name)
st.write("Report Summary:", report_summary)


# <h1>4. Prompt Chaining 1 - News Text to MITRE ATT&CK Technique</h1>

# ### 4.1 [Method 1] Using few-shot learning

# In[152]:


# prompt = f"""
# <Input_text>
# {scraped_text}
# </Input_text>

# <Instruction>
# I want you to act as a cyber security specialist. Based on the below provided examples, identify and extract the relevant \
# MITRE ATT&CK techniques from the input text provided above.

# <Examples>
# Few-shot Learning Examples:
# 1. "This file extracts credentials from LSASS similar to Mimikatz."
    # - MITRE ATT&CK Technique: T1003.001 (OS Credential Dumping: LSASS Memory)
    
# 2. "It spreads to Microsoft Windows machines using several propagation methods, including the EternalBlue exploit for the \
# CVE-2017-0144 vulnerability in the SMB service."
    # - MITRE ATT&CK Technique: T1210 (Exploitation of Remote Services)
    
# 3. "SMB copy and remote execution."
    # - MITRE ATT&CK Technique: T1570 (Lateral Tool Transfer)
# </Examples>

# Please provide the results in the following format:
# 1. "[Relevant text from the article]."
   # - MITRE ATT&CK Technique: [MITRE ATT&CK Technique Number] ([MITRE ATT&CK Technique Name])

# Make sure to structure the output as a numbered list, where each item consists of a short relevant text from the input \
# article followed by its corresponding MITRE ATT&CK technique and name.
# </Instruction>

# """

# response1 = get_completion(prompt)

# print(response1)


# #### Calculation of tokens

# In[153]:


#import tiktoken

#def estimate_token_counts(prompt, model='gpt-4o-mini'):
#    encoding = tiktoken.encoding_for_model(model)
#return len(encoding.encode(prompt))

#total_tokens_count = estimate_token_counts(prompt) + estimate_token_counts(response1)

#print(f"Total tokens count: {total_tokens_count}")


# ### 4.2 [Method 2] Using Retrieval-Augmented Generation (RAG)

# ##### Load and Embed the CSV for Retrieval

# In[154]:


# from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.document_loaders import CSVLoader
from langchain.vectorstores.faiss import FAISS
from langchain.schema import Document
from langchain.docstore.in_memory import InMemoryDocstore
from langchain_openai import OpenAIEmbeddings
import faiss

# Step 1: Load the CSV file using pandas
df_label = pd.read_csv('single_label.csv')

# Step 2: Convert the CSV rows into LangChain Document objects
docs = [
    Document(
        page_content=row['text'],
        metadata={
            'mitre_technique_ID': row['mitre_technique_ID'],
            'mitre_technique_name': row['mitre_technique_name']
        }
    ) for index, row in df_label.iterrows()
]

# Step 3: Create embeddings using OpenAIEmbeddings and pass the API key
embeddings = OpenAIEmbeddings(openai_api_key=openai_key)

# Step 4: Initialize FAISS index
index = faiss.IndexFlatL2(len(embeddings.embed_query(" ")))

# Step 5: Initialize FAISS with embeddings and other necessary parameters
vector_store = FAISS(embeddings, index, InMemoryDocstore(), {})

# Step 6: Add the created documents to the vector store
vector_store.add_documents(docs)

# Step 7: Create a retriever from the vector store
retriever = vector_store.as_retriever(search_kwargs={"k": 10})


# ##### Create the RAG Chain Using LangChain

# In[156]:


from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.chains import RetrievalQA

# Use the existing OpenAI setup to instantiate a new model for LangChain
llm = ChatOpenAI(openai_api_key=openai_key, model_name="gpt-4-turbo")  # Corrected to use ChatOpenAI

# Define system prompt template
system_prompt = """
<Input_text>
{scraped_text}
</Input_text>

<Instruction>
I want you to act as a cybersecurity specialist focused on attacker behavior. Identify and extract only the relevant \
MITRE ATT&CK techniques that represent actions, tools, or methods used by attackers, as described in the input text provided \
above. 

Please **exclude** any recommendations, defensive guidance, or security advice, including descriptions from credible sources \
such as Microsoft that suggest preventative or mitigation strategies (e.g., instructions to enable security features, block \
settings, or run endpoint detection software).

Format each entry as follows, strictly limiting to attacker actions:
1. "[Relevant text from the article]."
   - MITRE ATT&CK Technique: [MITRE ATT&CK Technique Number] ([MITRE ATT&CK Technique Name])

Ensure the output is structured as a numbered list, with each item containing a short relevant text from the input article \
followed by the corresponding MITRE ATT&CK technique and name, focusing only on attacker-specific techniques.
</Instruction>
"""

# Define the prompt template using the system prompt and the input
prompt_template = ChatPromptTemplate.from_messages([
    ("system", system_prompt),
])

# Create the retrieval-augmented chain
rag_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type="stuff",  # "stuff" means embedding everything into one prompt
    retriever=retriever,  # The retriever uses FAISS + embeddings from CSV data
    return_source_documents=False  # Return the source documents
)


# ##### Integrate Prompt into the RAG Process

# In[157]:


def generate_with_rag(scraped_text, context):
    prompt = system_prompt.format(
        scraped_text=scraped_text,
        context=context
    )
    
    return get_completion(prompt)

retrieved_context = rag_chain({"query": scraped_text})

response = generate_with_rag(scraped_text, retrieved_context["result"])


# ##### Post-processing to replace outdated techniques with relevant techniques

# In[168]:


prompt = f"""
<Input_text>
{response}
</Input_text>

<Instruction>
Please review the text above and identify any instances of deprecated MITRE ATT&CK technique IDs. \
Remove each deprecated technique and replace it with the most contextually appropriate, currently valid \
MITRE ATT&CK technique. Here is a list of deprecated techniques to replace:
- T1043
- T1066
- T1086

Maintain the following output format:

1. "[Relevant text]."
   - MITRE ATT&CK Technique: [MITRE ATT&CK Technique Number] ([MITRE ATT&CK Technique Name])

Preserve the exact original format, replacing only deprecated techniques with the most relevant alternative \
techniques.
</Instruction>
"""

response = get_completion(prompt)

st.markdown("""
### MITRE ATT&CK Mapping Results
The extracted threat intelligence text has been analyzed and mapped to relevant MITRE ATT&CK techniques. 

Each entry below represents a specific action or method identified in the text, along with its corresponding MITRE ATT&CK tactic and technique ID. This mapping aids in understanding the tactics and techniques attackers may use, providing a structured view of potential threats.
""")
st.write(response)


# ### 4.3 Extract list and dict of MITRE ATT&CK ID from prompt response

# In[169]:


def extract_unique_mitre_ids(response):
    # Use a regular expression to find all occurrences of MITRE ATT&CK technique IDs (e.g., T1003.001, T1071)
    mitre_ids = re.findall(r'T\d{4}(?:\.\d+)?', response)
    
    # Convert the list to a set to store only unique values and return as a list
    unique_mitre_ids = list(set(mitre_ids))
    
    return unique_mitre_ids

mitre_id_list = extract_unique_mitre_ids(response)

# #### Map MITRE Tactic to Technique using csv as a data source

# In[170]:


def extract_unique_mitre_ids(response):
    # Use a regular expression to find all occurrences of MITRE ATT&CK technique IDs (e.g., T1003.001, T1071)
    mitre_ids = re.findall(r'T\d{4}(?:\.\d+)?', response)
    
    # Convert the list to a set to store only unique values
    unique_mitre_ids = list(set(mitre_ids))
    
    # Create a dictionary to store technique_id, technique_name, tactic, and url
    mitre_dict = [{"tactic": "", "technique_name": "", "technique_id": mitre_id, "url": ""} for mitre_id in unique_mitre_ids]
    
    return mitre_dict

# Extract unique MITRE IDs into the dictionary
mitre_dict = extract_unique_mitre_ids(response)

# Load techniques data from CSV
df_techniques = pd.read_csv("MITRE-Techniques.csv")

# Create dictionaries to map IDs to technique name, tactic, and URL
id_to_tactics = dict(zip(df_techniques['ID'], df_techniques['tactics']))
id_to_technique_name = dict(zip(df_techniques['ID'], df_techniques['name']))
id_to_url = dict(zip(df_techniques['ID'], df_techniques['url']))

# Fill in the "technique_name", "tactic", and "url" fields in the mitre_dict
for entry in mitre_dict:
    technique_id = entry['technique_id']
    
    # If the technique_id exists in the id_to_tactics dictionary, update the tactic field
    if technique_id in id_to_tactics:
        entry['tactic'] = id_to_tactics[technique_id]
    
    # If the technique_id exists in the id_to_technique_name dictionary, update the technique_name field
    if technique_id in id_to_technique_name:
        entry['technique_name'] = id_to_technique_name[technique_id]
    
    # If the technique_id exists in the id_to_url dictionary, update the url field
    if technique_id in id_to_url:
        entry['url'] = id_to_url[technique_id]


# #### Convert dictionary into a dataframe and sort it according to MITRE tactic sequence

# In[171]:


# List of ordered tactics as per MITRE ATT&CK sequence
tactic_order = [
    "Reconnaissance", "Resource Development", "Initial Access", "Execution", "Persistence",
    "Privilege Escalation", "Defense Evasion", "Credential Access", "Discovery", "Lateral Movement",
    "Collection", "Command and Control", "Exfiltration", "Impact"
]

# Convert the mitre_dict into a DataFrame
df_mitre = pd.DataFrame(mitre_dict)

# Split the 'tactic' column by commas, explode the dataframe to create separate rows for each tactic
df_mitre['tactic'] = df_mitre['tactic'].str.split(', ')
df_mitre = df_mitre.explode('tactic')

# Create a categorical column for sorting based on tactic_order
df_mitre['tactic'] = pd.Categorical(df_mitre['tactic'], categories=tactic_order, ordered=True)

# Sort the DataFrame based on the tactic order
df_mitre = df_mitre.sort_values('tactic').reset_index(drop=True)

# Rename
df_mitre = df_mitre.rename(columns={
    'tactic': 'Tactic',
    'technique_name': 'Technique Name',
    'technique_id': 'Technique ID',
    'url': 'URL'
})

st.markdown("""
### Summary of Extracted Techniques
The table below provides a summary of the MITRE ATT&CK techniques and tactics identified in the analyzed text. Each row lists the tactic, associated technique, technique ID, and a link to MITRE's official documentation for further details on the technique.
""")

# Display the final sorted DataFrame
df_mitre


# #### Create technique_id to mitigation

# In[172]:


# Load the CSV into a pandas DataFrame
df_mitigations = pd.read_csv('MITRE-mitigations.csv')

# Filter df_mitigations to only rows where 'source ID' matches one of the 'mitre_id_list'
filtered_mitigations = df_mitigations[df_mitigations['target ID'].isin(mitre_id_list)]

# Initialize an empty dictionary to store mitigations by technique ID
mitigation_dict = {}

# Loop through each unique 'source ID' and aggregate the corresponding mitigations
for technique_id in mitre_id_list:
    # Get the matching mitigations for this technique
    matching_rows = filtered_mitigations[filtered_mitigations['target ID'] == technique_id]
    
    # Format the mitigations as requested
    mitigations = [f"[{row['source ID']}] {row['mapping description']}" for _, row in matching_rows.iterrows()]
    
    # Store the list of mitigations in the dictionary
    mitigation_dict[technique_id] = '\n\n'.join(mitigations)  # Join with double newlines for spacing between mitigations

# Convert the mitigation_dict to a DataFrame
df_mitigation_table = pd.DataFrame(list(mitigation_dict.items()), columns=['Technique ID', 'Mitigation'])

# Set the index of df_mitre to technique_id to ensure correct ordering
df_mitre_sorted = df_mitre.set_index('Technique ID')

# Reindex df_mitigation_table based on the sorted order of technique_id in df_mitre
df_mitigation_table = df_mitigation_table.set_index('Technique ID').reindex(df_mitre_sorted.index).reset_index()

# Drop duplicate rows based on the technique_id column
df_mitigation_table = df_mitigation_table.drop_duplicates(subset=['Technique ID'])

# Fill missing values in the 'mitigation' column with a standard string
default_message = "This type of attack technique cannot be easily mitigated with preventive controls since it is based on the abuse of system features."
df_mitigation_table['Mitigation'] = df_mitigation_table['Mitigation'].fillna(default_message).replace('', default_message)

st.markdown("""
### Recommended Mitigation Measures
This table presents recommended mitigation measures sourced from the MITRE ATT&CK framework. Each technique ID has been matched with specific actions or policies to help users efficiently implement security controls and minimize risks associated with these techniques.
""")

# Output the new DataFrame
df_mitigation_table


# <h1>5. For converting git clone Splunk .yml files into a dataframe</h1>

# ### Only run it if the section 1.2 code was ran, which git clones the .yml files

# In[163]:


# # Initialize the directory to where .yml files are located
# directory = r'C:\Users\Forensic\Desktop\AI LLM\security_content\detections\endpoint'

# # Initialize an empty list to store the extracted data
# data = []

# # Iterate through each .yml file in the directory
# for filename in os.listdir(directory):
#     if filename.endswith(".yml"):  # Only process .yml files
#         file_path = os.path.join(directory, filename)

#         # Open and load the .yml file
#         with open(file_path, 'r') as file:
#             try:
#                 # Load the content of the .yml file
#                 content = yaml.safe_load(file)
                
#                 # Extract the required fields
#                 name = content.get('name', 'N/A')  # Default to 'N/A' if key not found
#                 description = content.get('description', 'N/A')
#                 data_source = content.get('data_source', [])
                
#                 # Extract the 'tags' fields (confidence, impact, risk_score)
#                 tags = content.get('tags', {})
#                 mitre_attack_id = tags.get('mitre_attack_id', [])
                
#                 # Ensure the confidence, impact, and risk_score are treated as integers, defaulting to 0 if not present
#                 confidence = tags.get('confidence', 0)  # Default to 0 if missing
#                 impact = tags.get('impact', 0)  # Default to 0 if missing
#                 risk_score = tags.get('risk_score', 0)  # Default to 0 if missing
                
#                 # Extract the 'search' field as a string (which contains the Splunk query)
#                 search = content.get('search', 'N/A')
                
#                 # Append the extracted data to the list
#                 data.append({
#                     'name': name,
#                     'description': description,
#                     'data_source': ', '.join(data_source),  # Join list into a string
#                     'mitre_attack_id': ', '.join(mitre_attack_id),  # Join list into a string
#                     'search': search,
#                     'confidence': confidence,
#                     'impact': impact,
#                     'risk_score': risk_score
#                 })
#             except yaml.YAMLError as e:
#                 print(f"Error reading {filename}: {e}")

# # Convert the data list into a DataFrame
# df = pd.DataFrame(data)

# # Save the DataFrame to a CSV file
# df.to_csv('splunk_query_data.csv', index=False)

# # Display a message confirming the save
# print("Data has been saved to 'splunk_query_data.csv'")


# <h1>6. Prompt Chaining 2 - Splunk Query to KQL Query</h1>

# #### 1. Read Splunk query csv into a dataframe and create new DataFrame based on criteria of (i) 4688 data_source, (ii) confidence >50

# In[173]:


df_splunk = pd.read_csv("splunk_query_data.csv")

# Filter DataFrame where 'data_source' contains '4688' and 'confidence' is greater than 50
filtered_df = df_splunk[(df_splunk['data_source'].str.contains('4688', na=False)) & (df_splunk['confidence'] > 50)].reset_index(drop=True)


# #### 2. Filter out rows with mitre_attack_id match as those identified in threat intelligence

# In[174]:


def find_matching_mitre_ids(cell_value):
    if isinstance(cell_value, str):
        cell_mitre_ids = [mitre_id.strip() for mitre_id in cell_value.split(',')]
        matches = [mitre_id for mitre_id in cell_mitre_ids if mitre_id in mitre_id_list]
        return ', '.join(matches) if matches else None
    return None

# Apply the function and filter rows with matches
filtered_df['matching_mitre_id'] = filtered_df['mitre_attack_id'].apply(find_matching_mitre_ids)
final_filtered_df = filtered_df[filtered_df['matching_mitre_id'].notna()].reset_index(drop=True)

# Sort
final_filtered_df = final_filtered_df.sort_values(by='matching_mitre_id', ascending=True).reset_index(drop=True)


# #### 3. Transform splunk query into KQL Queries using few-shot learning

# In[175]:


# Initialize an empty list to store each numbered query
splunk_query_list = []

# Iterate over each row retrieving both 'matching_mitre_id' and 'search' columns
for index, row in enumerate(final_filtered_df.itertuples(), start=1):
    # Get the technique ID and the query
    mitre_id = row.matching_mitre_id
    query = row.search
    
    # Format each query with its corresponding MITRE ID and number
    formatted_query = f"[{mitre_id}]\n{index}. {query}"
    
    # Append the formatted query to the list
    splunk_query_list.append(formatted_query)

# Join the list into a single string with '\n\n' in between each query
splunk_queries = '\n\n'.join(splunk_query_list)

# splunk_queries contains the numbered search queries with MITRE IDs and newlines between them
print(splunk_queries)


# In[176]:


prompt = f"""
<Splunk queries>
{splunk_queries}
</Splunk queries>

<Instruction>
I want you to act as a cyber security specialist. Based on the examples provided below, transform the above Splunk query \
into its corresponding KQL query using the following rules:

1. Start each KQL query with the corresponding technique ID in brackets, exactly as it appears before each Splunk query \
(e.g., [T1003.001]). Ensure this `[Technique ID]` is preserved and placed at the beginning of each query output.
2. Start the KQL query with the correct log source, which is "SecurityEvent", and ensure that the KQL query always filters on \
"EventID == 4688".
3. Map the Splunk field names to KQL field names as follows:
   - `Processes.process_name` in Splunk should be mapped to `Process` in KQL (this represents the executable name, e.g., `cmd.exe`).
   - `Processes.process` in Splunk should be mapped to `CommandLine` in KQL (this represents the full command line, including arguments).
   - `Processes.process_path` in Splunk should be mapped to `NewProcessName` in KQL.
   - `Processes.dest` in Splunk should be mapped to `Computer` in KQL.
   - `Processes.user` in Splunk should be mapped to `Account` in KQL.
   - `Processes.parent_process_name` in Splunk should be mapped to `ParentProcessName` in KQL.
4. Ensure the case sensitivity of field names is preserved (e.g., `Process`, `CommandLine`, `Computer`).
5. For the `Process` field in KQL, always use `contains` and ensure that the process names include `.exe` (e.g., `Process contains "curl.exe"`).
6. Use `CommandLine` in KQL when searching for command-line arguments or specific strings in the command executed.
7. Exclude any aggregation functions like `summarize` or `count`. Ensure that the query outputs raw data logs for analysis.
8. Ensure to return the final list of queries in the correct KQL format only without any backticks or any comments.

<Examples>
### Example 1: Splunk Query to KQL Transformation

Splunk Query:
[T1003.001]
| tstats `security_content_summariesonly` count min(_time) as firstTime max(_time) as lastTime from datamodel=Endpoint.Processes where Processes.process_name=curl (Processes.process="*-s *") OR (Processes.process="*|*" AND Processes.process="*bash*") by Processes.dest Processes.user Processes.parent_process_name Processes.process_name Processes.process Processes.process_id Processes.parent_process_id

KQL Query:
[T1003.001]
SecurityEvent
| where EventID == 4688
| where Process contains "curl.exe" and (CommandLine contains "-s" or (CommandLine contains "|" and CommandLine contains "bash"))

### Example 2: Splunk Query to KQL Transformation

Splunk Query:
[T1059]
| tstats `security_content_summariesonly` values(Processes.process) as cmdline values(Processes.parent_process_name) as parent_process values(Processes.process_name) count min(_time) as firstTime max(_time) as lastTime from datamodel=Endpoint.Processes where (Processes.process_name = "netsh.exe" OR Processes.original_file_name= "netsh.exe") AND Processes.process = "*firewall*" AND Processes.process = "*add*" AND Processes.process = "*protocol=TCP*" AND Processes.process = "*localport=3389*" AND Processes.process = "*action=allow*" by Processes.dest Processes.user Processes.parent_process Processes.process_name Processes.process Processes.process_id Processes.parent_process_id | `drop_dm_object_name(Processes)` | `security_content_ctime(firstTime)` | `security_content_ctime(lastTime)` | `windows_remote_services_allow_rdp_in_firewall_filter`

KQL Query:
[T1059]
SecurityEvent
| where EventID == 4688
| where (Process contains "netsh.exe" and CommandLine contains "firewall" and CommandLine contains "add" and CommandLine contains "protocol=TCP" and CommandLine contains "localport=3389" and CommandLine contains "action=allow")

</Examples>

Make sure to structure the output as a numbered list, starting each item with the `[Technique ID]` exactly as shown, followed by the KQL query accurately transformed according to the rules provided above.
</Instruction>
"""

response_query = get_completion(prompt)


st.markdown("""
### Download Generated Files
The app has created three downloadable files, each serving a different purpose to support your cybersecurity analysis. Follow the steps below to understand and use each file effectively:

1. **`mitre_heatmap_layer.json`**: This file is ready to upload to the MITRE ATT&CK Navigator. You can visualize techniques and tactics by going to [MITRE ATT&CK Navigator](https://mitre-attack.github.io/attack-navigator/), selecting *'Open Existing Layer'* > *'Upload from local'*, and choosing this `.json` file.

2. **`KQL_query.txt`**: This text file contains a list of KQL queries designed for Microsoft Sentinel's Log Analytics Workspace. Each query corresponds to a specific MITRE ATT&CK technique, helping you to efficiently search for related threats in your environment. Simply copy and paste the queries as needed for your threat-hunting activities.

3. **`Final_report.docx`**: This Word document includes all the information presented in this app, along with detailed summaries and extracted intelligence. After generating a heatmap in MITRE ATT&CK Navigator, consider inserting it into this report for a comprehensive view of your findings.
""")

# <h1>7. Output - Report, JSON, KQL Query</h1>

# ### 7.1 JSON File For ATT&CK Navigator Heatmap

# In[177]:


# Load the existing JSON template
with open('heatmap_json_template.json', 'r') as json_file:
    json_data = json.load(json_file)

# Function to convert tactics to lowercase and replace spaces with dashes
def format_tactic(tactic):
    return tactic.lower().replace(' ', '-')

# Function to append techniques to the JSON
def append_techniques_to_json(techniques, json_data):
    for technique in techniques:
        tactics = technique['tactic'].split(', ')  # Split multiple tactics
        
        for tactic in tactics:
            json_data['techniques'].append({
                "techniqueID": technique['technique_id'],
                "tactic": format_tactic(tactic),  # Format the tactic
                "color": "#e60d0d",
                "comment": "-no comment-",
                "enabled": True,
                "metadata": [],
                "links": [],
                "showSubtechniques": False
            })
    
    return json_data

# Append the techniques to the JSON data
json_data = append_techniques_to_json(mitre_dict, json_data)

# Save the updated JSON data back to a file
with open('1. mitre_heatmap_layer.json', 'w') as output_file:
    json.dump(json_data, output_file, indent=2)

st.write("Saved to 1. mitre_heatmap_layer.json")

with open("1. mitre_heatmap_layer.json", "r") as file:
    json_content = file.read()
st.download_button(
    label="Download mitre_heatmap_layer.json",
    data=json_content,
    file_name="mitre_heatmap_layer.json",
    mime="application/json"
)

# ### 7.2 KQL Query Text File

# In[178]:


# Define the file path and name
file_name = "2. KQL_query.txt"

# Open the file in write mode and save the content of response_query
with open(file_name, 'w') as file:
    file.write(response_query)

st.write(f"Saved to {file_name}")

with open(file_name, "r") as file:
    kql_content = file.read()
st.download_button(
    label="Download KQL Query",
    data=kql_content,
    file_name=file_name,
    mime="text/plain"
)


# ### 7.3 Output Report

# #### 1. Initialise date time

# In[179]:



# #### 3. Final output document

# In[182]:


from docx import Document

# Load the document
doc = Document('Report_template.docx')

# Function to replace placeholder text in paragraphs
def replace_placeholder(paragraphs, placeholder, replacement_text):
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement_text)

# Function to add a DataFrame as a table to replace a specific placeholder in the document
def add_dataframe_table(doc, dataframe, placeholder):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Create the table with the same number of columns as the dataframe
            table = doc.add_table(rows=1, cols=len(dataframe.columns))

            # Add the header row with DataFrame column names
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(dataframe.columns):
                hdr_cells[i].text = col_name

            # Add rows from the DataFrame
            for _, row in dataframe.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Insert the table after the placeholder paragraph
            paragraph._element.addnext(table._tbl)
            
            # Remove the placeholder paragraph
            paragraph.text = ""
            break  # Exit after inserting the table to avoid affecting other placeholders

# Replace placeholders in text
replace_placeholder(doc.paragraphs, "{{export_time}}", export_time)
replace_placeholder(doc.paragraphs, "{{report_name}}", report_name)
replace_placeholder(doc.paragraphs, "{{report_summary}}", report_summary)
replace_placeholder(doc.paragraphs, "{{text_technique_mapping}}", response)      
            
# Replace {{table_technique}} with df_mitre
add_dataframe_table(doc, df_mitre, "{{table_technique}}")

# Replace {{table_mitigation}} with df_mitigation_table
add_dataframe_table(doc, df_mitigation_table, "{{table_mitigation}}")

# Save the modified document
doc.save('3. Final_Report.docx')

st.write("Saved to 3. Final_Report.docx")

with open("3. Final_Report.docx", "rb") as file:
    st.download_button(
        label="Download Final Report (DOCX)",
        data=file,
        file_name="Final_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# <h1>8. Model Evaluation</h1>

# In[30]:


# from datasets import Dataset
# from ragas import evaluate
# from ragas.metrics import faithfulness, context_recall, context_precision, answer_relevancy

# # Set the API key as an environment variable for other libraries
# os.environ["OPENAI_API_KEY"] = openai_key  # This ensures other libraries can access the key

# # Sample scraped text from URL or PDF
# scraped_text = "Advanced Persistent Threat (APT) actors delivered spear-phishing emails leading to the execution of T1059.001 and T1071."

# # Example generated output (this would be the output from RAG pipeline)
# generated_ttp_ids = ["T1059.001", "T1071"]

# # Reference or ground truth TTPs (the expected correct TTPs for comparison)
# ground_truth_ttps = ["T1059.001", "T1071", "T1021"]

# # Add the scraped text into the contexts list
# list_of_contexts = [scraped_text]

# # Prepare the dataset for evaluation
# data_samples = {
#     'question': ['Extract relevant MITRE ATT&CK TTPs from the input text.'],
#     'answer': [", ".join(generated_ttp_ids)],  # Join the list of generated TTPs into a single string
#     'contexts' : [list_of_contexts],  # Context is the scraped text
#     'ground_truth': [", ".join(ground_truth_ttps)]  # Join the ground truth TTPs into a single string
# }

# # Convert the dictionary into a Dataset object
# dataset = Dataset.from_dict(data_samples)

# # Evaluate the dataset using the ragas metrics
# score = evaluate(dataset, metrics=[faithfulness, context_recall, context_precision, answer_relevancy])

# # Convert the score to a pandas DataFrame for better readability
# score_df = score.to_pandas()

# # Output the score
# score_df

